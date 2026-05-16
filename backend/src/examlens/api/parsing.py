from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path


@dataclass
class ParsedDocument:
    text: str
    title: str | None = None
    pages: list[str] | None = None
    headings: list[dict[str, Any]] | None = None


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"-\n(?=\w)", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_headings(text: str) -> list[dict[str, Any]]:
    """Simple heading extraction based on common patterns."""
    headings = []
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line = line.strip()
        # Matches "Chapter 1", "1. Introduction", "## Heading"
        if re.match(r"^(#{1,3}\s+|Chapter\s+\d+|[A-Z\d]+\.\s+[A-Z])", line, re.I):
            headings.append({"text": line, "line": i})
    return headings


def parse_pdf(path: Path) -> ParsedDocument:
    import fitz

    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(_clean_text(page.get_text("text")))
    full_text = "\n\n".join(pages)
    return ParsedDocument(
        text=full_text,
        title=path.stem,
        pages=pages,
        headings=_extract_headings(full_text)
    )


def parse_docx(path: Path) -> ParsedDocument:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    text = _clean_text("\n".join(parts))
    return ParsedDocument(text=text, title=path.stem)


def parse_text_file(path: Path) -> ParsedDocument:
    text = _clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    return ParsedDocument(text=text, title=path.stem)


def parse_markdown_file(path: Path) -> ParsedDocument:
    return parse_text_file(path)


def parse_any(path: Path, mime_type: str) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf" or mime_type == "application/pdf":
        return parse_pdf(path)
    if suffix in {".docx"}:
        return parse_docx(path)
    if suffix in {".md", ".markdown"}:
        return parse_markdown_file(path)
    return parse_text_file(path)


def checksum_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()

